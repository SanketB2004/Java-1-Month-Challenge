from docx import Document

# Create Word document
doc = Document()
doc.add_heading('Resume (बायोडेटा)', level=1)

doc.add_heading('व्यक्तिगत माहिती / Personal Details', level=2)
doc.add_paragraph('नाव: [तुमचं पूर्ण नाव]')
doc.add_paragraph('पत्ता: [तुमचा संपूर्ण पत्ता]')
doc.add_paragraph('फोन नंबर: +91‑XXXXXXXXXX')
doc.add_paragraph('ई‑मेल: yourname@email.com')
doc.add_paragraph('GitHub: [GitHub प्रोफाइल लिंक]')
doc.add_paragraph('LinkedIn: [LinkedIn प्रोफाइल लिंक]')

doc.add_heading('करिअर उद्दिष्ट / Career Objective', level=2)
doc.add_paragraph('मी एक मेहनती आणि जिज्ञासू विद्यार्थी आहे...')

doc.add_heading('शैक्षणिक पात्रता / Educational Qualification', level=2)
table = doc.add_table(rows=1, cols=4)
hdr = table.rows[0].cells
hdr[0].text = 'पदवी / परीक्षा'
hdr[1].text = 'संस्था / बोर्ड'
hdr[2].text = 'वर्ष'
hdr[3].text = 'गुण (%) / CGPA'
rows = [
    ['B.Sc. / BCA / BE (CSE/IT)', '[कॉलेजचं नाव]', '2023', '75%'],
    ['HSC (१२वी)', '[बोर्ड नाव]', '2020', '80%'],
    ['SSC (१०वी)', '[बोर्ड नाव]', '2018', '85%']
]
for r in rows:
    row_cells = table.add_row().cells
    for i, val in enumerate(r):
        row_cells[i].text = val

doc.add_heading('Technical Skills', level=2)
skills = [
    'Languages: Java, C, C++, Python',
    'Web: HTML, CSS, JavaScript',
    'Database: MySQL, Oracle',
    'Tools & Platforms: Git, GitHub, Eclipse, VS Code',
    'OS: Windows, Linux'
]
for s in skills:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Projects', level=2)
projects = [
    'Student Information System – Java, MySQL, JDBC (GitHub link)',
    'Online Voting System – HTML, CSS, JavaScript, PHP (GitHub link)'
]
for p in projects:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('Certifications', level=2)
certs = [
    'Java Programming – Coursera / Udemy',
    'Python for Beginners – Great Learning',
    'Web Development Bootcamp – Internshala / Udemy',
    'SQL Fundamentals – HackerRank'
]
for c in certs:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Achievements', level=2)
achs = [
    'Participated in Smart India Hackathon – 2023',
    '2nd prize in College-Level Coding Contest',
    'Solved 100+ problems on LeetCode / HackerRank',
    'Led TechFest 2023 in college'
]
for a in achs:
    doc.add_paragraph(a, style='List Bullet')

doc.add_heading('भाषा कौशल्ये / Languages Known', level=2)
langs = ['मराठी – मातृभाषा', 'हिंदी – उत्तम', 'इंग्रजी – उत्तम']
for l in langs:
    doc.add_paragraph(l, style='List Bullet')

doc.add_paragraph('\nघोषणा:\nवरील माहिती खरी आणि अचूक आहे...')
doc.add_paragraph('दिनांक: __________\nस्वाक्षरी: __________')

# Save .docx file
doc.save('Marathi_Resume_Template.docx')
